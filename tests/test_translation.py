import os
import unittest
import docx
from backend.services.language_detection_service import LanguageDetectionService
from backend.services.translation_service import translation_engine
from backend.services.pdf_generation_service import pdf_generator

class TestUniversalTranslationService(unittest.TestCase):

    def test_multi_language_detection(self):
        sample_texts = {
            "Marathi": "महाराष्ट्र शासन परिपत्रक क्रमांक 1021 जिल्हाधिकारी कार्यालय",
            "Hindi": "उत्तर प्रदेश सरकार अधिसूचना संख्या 405 राजस्व विभाग",
            "Spanish": "Este es un documento oficial del gobierno de España.",
            "French": "Ceci est un document officiel du gouvernement français.",
            "English": "Official Gazette Election Notification Polling Station"
        }
        for expected_lang, text in sample_texts.items():
            detected = LanguageDetectionService.detect_language(text)
            self.assertIsNotNone(detected)
            self.assertNotEqual(detected, "")

    def test_multi_language_translation_to_english(self):
        sample_texts = [
            ("Hindi", "नमस्ते दुनिया", "Hello"),
            ("Spanish", "Este es un documento oficial.", "official document"),
            ("French", "Gouvernement français", "French government"),
            ("Marathi", "महाराष्ट्र शासन", "Government of Maharashtra")
        ]
        for src_lang, text, expected_substr in sample_texts:
            translated = translation_engine.translate_paragraph(text, src_lang)
            self.assertIsNotNone(translated)
            self.assertTrue(len(translated) > 0)

    def test_docx_in_place_format_preservation(self):
        # Create a sample DOCX file with headings, paragraphs, and tables
        sample_docx_path = "tests/temp_sample_in.docx"
        output_docx_path = "tests/temp_sample_out.docx"
        
        doc = docx.Document()
        doc.add_heading("दस्तावेज़ शीर्षक", level=1)
        doc.add_paragraph("यह पहला पैराग्राफ है।")
        
        t = doc.add_table(rows=1, cols=2)
        t.cell(0, 0).text = "राज्य"
        t.cell(0, 1).text = "महाराष्ट्र"
        doc.save(sample_docx_path)

        # Run generate_translated_docx with in-place original_path
        res_path = pdf_generator.generate_translated_docx(
            doc_id="test1234",
            filename="temp_sample_in.docx",
            paragraphs=[{"translated_text": "Translated"}],
            metadata={"subject": "Test Doc"},
            original_path=sample_docx_path
        )
        self.assertTrue(os.path.exists(res_path))

        out_doc = docx.Document(res_path)
        self.assertTrue(len(out_doc.paragraphs) > 0)
        self.assertTrue(len(out_doc.tables) > 0)

        # Clean up
        if os.path.exists(sample_docx_path): os.remove(sample_docx_path)
        if os.path.exists(res_path): os.remove(res_path)

    def test_marathi_proverbs_with_ocr_noise(self):
        test_cases = [
            ("१. अति तिथे मातिी", "1. Excess of anything is harmful."),
            ("२. पेरावे तिसे उगवतिे", "2. As you sow, so shall you reap."),
            ("३. उथळ पाण्याला खळखळाट जास्ति", "3. Empty vessels make the most noise."),
            ("४. आयत्या बिळावर नागोबिा", "4. Taking advantage of another's effort."),
            ("५. दुरून डोंगर साजरे", "5. The grass is always greener on the other side.")
        ]
        for raw_text, expected_translation in test_cases:
            res = translation_engine.translate_paragraph(raw_text, "Marathi")
            self.assertEqual(res, expected_translation)

    def test_marathi_honorific_affirmative_translation(self):
        phrase = "ज्येष्ठ व गुणवंत कलावंतांचा गौरव व सन्मान केला"
        res = translation_engine.translate_paragraph(phrase, "Marathi")
        self.assertNotIn("not", res.lower())
        self.assertIn("felicita", res.lower())
        self.assertIn("honor", res.lower())

    def test_nabard_energy_irrigation_translation(self):
        corrupted_input = "ऊजार्ट बचति के साथ नबार्टधि सर्तिचाई सुविधिा उपलब्ध कराना"
        res = translation_engine.translate_paragraph(corrupted_input, "Marathi")
        self.assertNotIn("nabartdhi", res.lower())
        self.assertNotIn("srtichhai", res.lower())
        self.assertNotIn("uzart", res.lower())
        self.assertIn("nabard", res.lower())
        self.assertIn("irrigation", res.lower())

    def test_taleghar_proper_noun_translation(self):
        place_input = "ग्रामपंचायत तळेघर"
        res = translation_engine.translate_paragraph(place_input, "Marathi")
        self.assertNotIn("basement", res.lower())
        self.assertIn("taleghar", res.lower())

if __name__ == "__main__":
    unittest.main()
