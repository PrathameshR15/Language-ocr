import asyncio
import os
import docx
import pytest
from backend.services.document_parser_service import document_parser

@pytest.mark.asyncio
async def test_e2e_flow():
    print("--- Testing DOCX Upload & Translation ---")
    docx_filename = "e2e_test_hindi.docx"
    doc = docx.Document()
    doc.add_heading("महाराष्ट्र राज्य अधिसूचना", level=1)
    p = doc.add_paragraph("यह एक आधिकारिक दस्तावेज है।")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "कार्यालय"
    table.cell(0, 1).text = "जिल्हाधिकारी कार्यालय मुंबई"
    doc.save(docx_filename)

    with open(docx_filename, "rb") as f:
        docx_bytes = f.read()

    res_docx = await document_parser.process_document(docx_bytes, docx_filename)
    print("DOCX Result ID:", res_docx["id"])
    print("DOCX Language Detected:", res_docx["language"])
    print("DOCX Paragraphs extracted:", len(res_docx["paragraphs"]))
    assert res_docx["original_pdf_path"] != ""
    assert res_docx["translated_pdf_path"] == ""
    assert len(res_docx["paragraphs"]) > 0
    print("First Para Text length:", len(res_docx["paragraphs"][0]["text"]))
    print("First Para Translation length:", len(res_docx["paragraphs"][0]["translated_text"]))

    print("\n--- Testing TXT Upload & Translation ---")
    txt_filename = "e2e_test_bengali.txt"
    with open(txt_filename, "w", encoding="utf-8") as f:
        f.write("পশ্চিমবঙ্গ সরকার পরিবেশ ও শিক্ষা বিভাগ।\nতারিখ: ২৯ জুলাই ২০২৬।")

    with open(txt_filename, "rb") as f:
        txt_bytes = f.read()

    res_txt = await document_parser.process_document(txt_bytes, txt_filename)
    print("TXT Result ID:", res_txt["id"])
    print("TXT Language Detected:", res_txt["language"])
    print("TXT Paragraphs extracted:", len(res_txt["paragraphs"]))
    assert res_txt["original_pdf_path"] != ""
    assert res_txt["translated_pdf_path"] == ""
    assert len(res_txt["paragraphs"]) > 0

    if os.path.exists(docx_filename): os.remove(docx_filename)
    if os.path.exists(txt_filename): os.remove(txt_filename)

    print("\n[SUCCESS] E2E Verification Complete & Successful!")

if __name__ == "__main__":
    asyncio.run(test_e2e_flow())
