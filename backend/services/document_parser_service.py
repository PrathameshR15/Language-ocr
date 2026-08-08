import os
import re
import uuid
import time
import json
import tempfile
# pyrefly: ignore [missing-import]
import fitz  # PyMuPDF
# pyrefly: ignore [missing-import]
from PIL import Image
# pyrefly: ignore [missing-import]
from docx import Document as DocxDocument
from concurrent.futures import ThreadPoolExecutor
from typing import Dict, Any, List

from config import settings
from backend.utils.logger import logger
from backend.utils.file_validator import sanitize_filename, validate_file_size, validate_file_extension, detect_corrupted_file
from backend.services.ocr_service import ocr_engine
from backend.services.language_detection_service import LanguageDetectionService
from backend.services.translation_service import translation_engine, INDIC_ADMIN_DICTIONARY, MARATHI_IDIOM_GLOSSARY, is_corrupted_romanized_marathi
from backend.services.metadata_service import metadata_service
from backend.services.pdf_generation_service import pdf_generator
from backend.services.quality_validation_service import quality_validator
from backend.services.llm_enhancement_service import llm_service
from backend.database.file_db import db_client

class DocumentParserService:
    """Master Pipeline Service executing Steps 1 to 13 of document translation workflow."""

    @staticmethod
    async def process_document(file_bytes: bytes, original_filename: str) -> Dict[str, Any]:
        start_time = time.time()
        doc_id = str(uuid.uuid4())[:8]
        clean_name = sanitize_filename(original_filename)

        # Clear in-memory translation cache to guarantee fresh translation pipeline execution
        translation_engine.clear_cache()

        # Step 2: Validate
        validate_file_size(file_bytes)
        ext = validate_file_extension(clean_name)


        paragraphs: List[Dict[str, Any]] = []
        overall_confidence = 0.90
        doc_type = "Image"

        # Use temporary directory for transient file reading during parsing (auto-deleted)
        with tempfile.TemporaryDirectory() as tmpdir:
            temp_input_path = os.path.join(tmpdir, clean_name)
            with open(temp_input_path, "wb") as f:
                f.write(file_bytes)

            # Detect corruption
            detect_corrupted_file(temp_input_path, ext)

            # Step 3, 4, 5: Determine type & extract paragraphs using Hybrid Classifier
            if ext == "pdf":
                digital_paras, is_digital = ocr_engine.extract_digital_pdf_blocks(file_bytes)
                if is_digital:
                    is_valid, reason = ocr_engine.validate_ocr_quality(digital_paras, avg_conf=1.00)
                    if not is_valid:
                        logger.warning(f"Doc {doc_id}: Selectable digital text failed quality check: {reason}. Forcing scanned PDF OCR fallback...")
                        is_digital = False

                if is_digital:
                    doc_type = "Digital PDF"
                    paragraphs = digital_paras
                    overall_confidence = 1.00
                    logger.info(f"Doc {doc_id} classified as Digital PDF. Extracted {len(paragraphs)} digital blocks (OCR bypassed).")
                else:
                    doc_type = "Scanned PDF"
                    logger.info(f"Doc {doc_id} identified as Scanned PDF/Corrupted Font CMap. Processing page images with OCR...")
                    from backend.services.ocr_service import sanitize_indic_text, _ensure_tesseract
                    # pyrefly: ignore [missing-import]
                    from PIL import ImageOps, ImageDraw
                    # pyrefly: ignore [missing-import]
                    import pytesseract
                    
                    pdf_doc = fitz.open(temp_input_path)
                    
                    # Group digital table blocks by page so we can reconstruct them
                    digital_tables_by_page = {}
                    if digital_paras:
                        for p in digital_paras:
                            if p.get("block_type") == "table" and p.get("table_grid"):
                                pg = p.get("page", 1)
                                digital_tables_by_page.setdefault(pg, []).append(p)
                    
                    for page_num, page in enumerate(pdf_doc, start=1):
                        pix = page.get_pixmap(dpi=200)
                        img_path = os.path.join(tmpdir, f"temp_{doc_id}_p{page_num}.png")
                        pix.save(img_path)
                        
                        with Image.open(img_path) as pil_img:
                            page_w = page.rect.width
                            page_h = page.rect.height
                            
                            # Scale factor from PDF points to full-res image pixels
                            img_w, img_h = pil_img.size
                            scale_x = img_w / page_w
                            scale_y = img_h / page_h
                            
                            page_tables = digital_tables_by_page.get(page_num, [])
                            
                            # 1. OCR table cells individually using PyMuPDF cell coordinates
                            for tbl in page_tables:
                                grid = tbl.get("table_grid", [])
                                cell_bboxes = tbl.get("table_cell_bboxes", [])
                                
                                num_rows = len(grid)
                                num_cols = max(len(r) for r in grid) if grid else 0
                                new_grid = [["" for _ in range(num_cols)] for _ in range(num_rows)]
                                
                                # Setup tesseract env
                                abs_tessdata_dir = os.path.abspath(settings.TESSERACT_DATA_DIR)
                                os.environ["TESSDATA_PREFIX"] = abs_tessdata_dir
                                
                                tess_lang = "ben"
                                try:
                                    tess_lang = ocr_engine.detect_languages_for_tesseract(pil_img, abs_tessdata_dir)
                                except Exception:
                                    pass
                                
                                for r_idx in range(num_rows):
                                    for c_idx in range(len(grid[r_idx])):
                                        if cell_bboxes and r_idx < len(cell_bboxes) and c_idx < len(cell_bboxes[r_idx]):
                                            c_bbox = cell_bboxes[r_idx][c_idx]
                                            if c_bbox and len(c_bbox) == 4:
                                                cx0 = int(c_bbox[0] * scale_x)
                                                cy0 = int(c_bbox[1] * scale_y)
                                                cx1 = int(c_bbox[2] * scale_x)
                                                cy1 = int(c_bbox[3] * scale_y)
                                                
                                                cx0 = max(0, cx0)
                                                cy0 = max(0, cy0)
                                                cx1 = min(img_w, cx1)
                                                cy1 = min(img_h, cy1)
                                                
                                                if cx1 > cx0 and cy1 > cy0:
                                                    cell_img = pil_img.crop((cx0, cy0, cx1, cy1))
                                                    cell_img = ImageOps.expand(cell_img, border=10, fill="white")
                                                    
                                                    cell_text = pytesseract.image_to_string(
                                                        cell_img,
                                                        lang=tess_lang,
                                                        config="--psm 6"
                                                    ).strip()
                                                    new_grid[r_idx][c_idx] = sanitize_indic_text(cell_text)
                                
                                tbl["table_grid"] = new_grid
                                paragraphs.append(tbl)
                            
                            # 2. Redact table regions from the main image to prevent duplicate OCR text
                            redacted_img = pil_img.copy()
                            draw = ImageDraw.Draw(redacted_img)
                            for tbl in page_tables:
                                t_bbox = tbl.get("bbox")
                                if t_bbox and len(t_bbox) == 4:
                                    tx0 = int(t_bbox[0] * scale_x)
                                    ty0 = int(t_bbox[1] * scale_y)
                                    tx1 = int(t_bbox[2] * scale_x)
                                    ty1 = int(t_bbox[3] * scale_y)
                                    draw.rectangle([tx0 - 5, ty0 - 5, tx1 + 5, ty1 + 5], fill="white")
                            
                            # 3. OCR the redacted image to get clean non-table paragraphs
                            page_paras, page_conf = ocr_engine.process_image(redacted_img, page_num)
                            
                            # Determine scaling between whole-page OCR image and PDF points
                            w, h = redacted_img.size
                            new_w, new_h = w, h
                            if max(w, h) > 1600:
                                scale = 1600.0 / float(max(w, h))
                                new_w, new_h = int(w * scale), int(h * scale)
                            
                            x_scale = page_w / float(new_w) if new_w > 0 else 1.0
                            y_scale = page_h / float(new_h) if new_h > 0 else 1.0
                            
                            for op in page_paras:
                                o_bbox = op.get("bbox")
                                if o_bbox and len(o_bbox) == 4:
                                    op["bbox"] = [
                                        o_bbox[0] * x_scale,
                                        o_bbox[1] * y_scale,
                                        o_bbox[2] * x_scale,
                                        o_bbox[3] * y_scale
                                    ]
                                paragraphs.append(op)
                    pdf_doc.close()

            elif ext in ["png", "jpeg", "jpg", "tiff", "bmp"]:
                doc_type = "Image"
                logger.info(f"Doc {doc_id} identified as Image document.")
                with Image.open(temp_input_path) as pil_img:
                    paragraphs, overall_confidence = ocr_engine.process_image(pil_img, page_num=1)

            elif ext == "docx":
                doc_type = "Word Document"
                logger.info(f"Doc {doc_id} identified as Word DOCX.")
                docx_obj = DocxDocument(temp_input_path)
                para_idx = 1
                
                # 1. Extract paragraphs
                for p in docx_obj.paragraphs:
                    txt = p.text.strip()
                    if txt:
                        paragraphs.append({
                            "paragraph": para_idx,
                            "page": 1,
                            "text": txt,
                            "confidence": 0.99
                        })
                        para_idx += 1

                # 2. Extract tables as whole table grid blocks
                for table in docx_obj.tables:
                    raw_grid = []
                    for row in table.rows:
                        row_cells = []
                        for cell in row.cells:
                            c_text = cell.text.strip()
                            if not row_cells or c_text != row_cells[-1]:
                                row_cells.append(c_text)
                        if any(row_cells):
                            raw_grid.append(row_cells)
                    if raw_grid:
                        orig_md_lines = ["| " + " | ".join(raw_grid[0]) + " |", "| " + " | ".join(["---"] * len(raw_grid[0])) + " |"]
                        for r in raw_grid[1:]:
                            orig_md_lines.append("| " + " | ".join(r) + " |")
                        paragraphs.append({
                            "paragraph": para_idx,
                            "page": 1,
                            "text": "\n".join(orig_md_lines),
                            "confidence": 0.99,
                            "block_type": "table",
                            "table_grid": raw_grid
                        })
                        para_idx += 1

            elif ext == "txt":
                doc_type = "Text Document"
                logger.info(f"Doc {doc_id} identified as Text File.")
                try:
                    with open(temp_input_path, "r", encoding="utf-8", errors="ignore") as tf:
                        content = tf.read().strip()
                    if content:
                        blocks = re.split(r'\n\s*\n', content)
                        para_idx = 1
                        for block in blocks:
                            cleaned_block = " ".join([line.strip() for line in block.splitlines() if line.strip()])
                            if cleaned_block:
                                paragraphs.append({
                                    "paragraph": para_idx,
                                    "page": 1,
                                    "text": cleaned_block,
                                    "confidence": 1.00
                                })
                                para_idx += 1
                except Exception as e:
                    logger.error(f"Error reading TXT file {temp_input_path}: {e}")

        # Fallback if no text extracted
        if not paragraphs:
            paragraphs = [{
                "paragraph": 1,
                "page": 1,
                "text": "No readable text detected in document.",
                "confidence": 0.50
            }]

        # Step 7.5: Post-OCR LLM Contextual Typo & Corruption Correction
        enable_llm_correction = getattr(settings, "ENABLE_LLM_OCR_CORRECTION", True)
        if enable_llm_correction and paragraphs and llm_service.is_available():
            try:
                logger.info(f"Doc {doc_id}: Running post-OCR LLM text correction on {len(paragraphs)} paragraphs...")
                paragraphs = llm_service.correct_ocr_paragraphs_with_llm(paragraphs)
            except Exception as e:
                logger.warning(f"Doc {doc_id}: Post-OCR LLM text correction skipped: {e}")

        # Step 8 & 9: Paragraph-wise Language Detection & Translation (Parallelized)
        from backend.services.translation_service import INDIC_ADMIN_DICTIONARY, normalize_indic_digits

        def _process_single_para(p):
            is_table_block = (p.get("block_type") == "table") or bool(p.get("table_grid"))
            if is_table_block and p.get("table_grid"):
                raw_grid = p.get("table_grid", [])
                
                # Pre-detect the language from the first non-empty cell value
                sample_text = ""
                for row in raw_grid:
                    for cell in row:
                        c_str = str(cell or "").strip()
                        if c_str and not re.match(r'^[0-9\s,\.\-\+\(\)₹%\/०-९]+$', c_str):
                            sample_text = c_str
                            break
                    if sample_text:
                        break
                
                p_lang, lang_conf = LanguageDetectionService.detect_language_with_confidence(sample_text or "Marathi")
                
                trans_grid = []
                for row_idx, row in enumerate(raw_grid):
                    trans_row = []
                    for col_idx, cell in enumerate(row):
                        c_str = str(cell or "").strip()
                        if c_str:
                            # 0. Table Serial Number Normalization:
                            # If we are in the first column, not the header, and cell is a known corrupted number/character
                            if col_idx == 0 and row_idx > 0 and (c_str in {"ते", "তে", "ডে", "ডে", "०", "০", "রে", "ре", "রে", ".", "২", "२"} or re.match(r'^[तेतेडेডে০०রে.]+$', c_str)):
                                t_cell = str(row_idx)
                            # 1. Exact administrative dictionary match
                            elif c_str in INDIC_ADMIN_DICTIONARY:
                                t_cell = INDIC_ADMIN_DICTIONARY[c_str]
                            # 2. Idiom / phrase glossary match
                            elif c_str in MARATHI_IDIOM_GLOSSARY:
                                t_cell = MARATHI_IDIOM_GLOSSARY[c_str]
                            # 3. Detect if corrupted Romanized Marathi or Hinglish (e.g. (Pine Ke Pani Ki Nai Pipeline))
                            # prior to checking English regex (Rule 3)
                            elif is_corrupted_romanized_marathi(c_str, p_lang):
                                t_cell = translation_engine.translate_paragraph(c_str, p_lang)
                            # 4. Fast optimization: If cell is numeric/currency/date string, normalize digits locally
                            elif re.match(r'^[0-9\s,\.\-\+\(\)₹%\/०-९]+$', c_str):
                                t_cell = normalize_indic_digits(c_str)
                            # 5. Already English text (which doesn't match romanized marathi/hinglish check above)
                            elif re.match(r'^[a-zA-Z0-9\s\.,\-\/\:\;\(\)₹%]+$', c_str):
                                t_cell = c_str
                            else:
                                t_cell = translation_engine.translate_paragraph(c_str, p_lang)
                            trans_row.append(t_cell)
                        else:
                            trans_row.append("")
                    trans_grid.append(trans_row)

                p["language"] = p_lang
                p["language_confidence"] = lang_conf
                p["translated_table_grid"] = trans_grid
                p["block_type"] = "table"

                # Construct clean Markdown tables for original and translated content
                orig_md_lines = []
                if raw_grid:
                    orig_md_lines.append("| " + " | ".join(raw_grid[0]) + " |")
                    orig_md_lines.append("| " + " | ".join(["---"] * len(raw_grid[0])) + " |")
                    for r in raw_grid[1:]:
                        orig_md_lines.append("| " + " | ".join(r) + " |")
                orig_txt = "\n".join(orig_md_lines) if orig_md_lines else p.get("text", "")
                p["text"] = orig_txt

                trans_md_lines = []
                if trans_grid:
                    trans_md_lines.append("| " + " | ".join(trans_grid[0]) + " |")
                    trans_md_lines.append("| " + " | ".join(["---"] * len(trans_grid[0])) + " |")
                    for r in trans_grid[1:]:
                        trans_md_lines.append("| " + " | ".join(r) + " |")
                trans_txt = "\n".join(trans_md_lines) if trans_md_lines else "[TRANSLATED TABLE GRID]"
                p["translated_text"] = trans_txt
                p["translated_table_markdown"] = trans_txt

                return p, orig_txt, trans_txt, p_lang
            else:
                orig_txt = p.get("text", "")
                p_lang, lang_conf = LanguageDetectionService.detect_language_with_confidence(orig_txt)
                p["language"] = p_lang
                p["language_confidence"] = lang_conf

                # Fast bypass if paragraph is already pure English
                if p_lang == "English" and re.match(r'^[a-zA-Z0-9\s\.,\-\/\:\;\(\)₹%]+$', orig_txt):
                    trans_txt = orig_txt
                elif orig_txt in INDIC_ADMIN_DICTIONARY:
                    trans_txt = INDIC_ADMIN_DICTIONARY[orig_txt]
                elif orig_txt in MARATHI_IDIOM_GLOSSARY:
                    trans_txt = MARATHI_IDIOM_GLOSSARY[orig_txt]
                else:
                    trans_txt = translation_engine.translate_paragraph(orig_txt, p_lang)

                p["translated_text"] = trans_txt
                return p, orig_txt, trans_txt, p_lang

        detected_languages = []
        full_original_text_list = []
        full_translated_text_list = []

        if paragraphs:
            max_workers = min(10, max(1, len(paragraphs)))
            with ThreadPoolExecutor(max_workers=max_workers) as executor:
                processed_items = list(executor.map(_process_single_para, paragraphs))

            for p, orig_txt, trans_txt, p_lang in processed_items:
                detected_languages.append(p_lang)
                full_original_text_list.append(orig_txt)
                full_translated_text_list.append(trans_txt)

        # Primary document language
        primary_lang = max(set(detected_languages), key=detected_languages.count) if detected_languages else "English"
        full_original_text = "\n\n".join(full_original_text_list)
        full_translated_text = "\n\n".join(full_translated_text_list)

        # Step 12: Generate Metadata
        metadata_json = metadata_service.extract_metadata(full_original_text + "\n" + full_translated_text, clean_name)
        metadata_json["document_type"] = doc_type

        # Smart filename resolution if uploaded filename is generic/random
        display_name = clean_name
        is_generic = (
            clean_name.lower().startswith(("government_document", "upload", "file", "tmp", "blob", "img_")) or
            re.match(r'^[a-f0-9\-_]{8,}$', clean_name.rsplit('.', 1)[0], re.IGNORECASE) or
            re.match(r'^\d+[\-_]', clean_name)
        )
        if is_generic:
            subj = metadata_json.get("subject", "")
            dept = metadata_json.get("department", "")
            doc_no = metadata_json.get("doc_number", "")
            
            if subj and subj != "N/A" and len(subj.strip()) > 3:
                clean_subj = sanitize_filename(subj)[:40]
                display_name = f"{clean_subj}.{ext}"
            elif dept and dept != "N/A" and doc_no and doc_no != "N/A":
                clean_dept = sanitize_filename(dept)[:25]
                clean_doc = sanitize_filename(doc_no)[:15]
                display_name = f"{clean_dept}_{clean_doc}.{ext}"
            elif clean_name.lower().startswith("government_document"):
                display_name = f"Document_{doc_id}.{ext}"

        processing_time = round(time.time() - start_time, 2)
        upload_time_str = time.strftime("%Y-%m-%d %H:%M:%S")

        # Save original uploaded document file to UPLOAD_DIR for layout-preserved in-place redaction / PDF preview
        orig_file_path = os.path.join(settings.UPLOAD_DIR, f"{doc_id}_{clean_name}")
        try:
            with open(orig_file_path, "wb") as f:
                f.write(file_bytes)
        except Exception as e:
            logger.warning(f"Could not save original upload file to disk: {e}")
            orig_file_path = ""

        metadata_json["original_pdf_path"] = orig_file_path

        # Step 9 & 10: Run Quality Validation & Calculate Multi-Tier Confidence Scores
        qa_results = quality_validator.validate_translation(paragraphs, metadata_json)
        
        ocr_conf = round(qa_results.get("ocr_accuracy", 98.0) / 100.0, 2)
        trans_conf = round(qa_results.get("translation_accuracy", 98.0) / 100.0, 2)
        layout_conf = round(qa_results.get("layout_similarity", 97.0) / 100.0, 2)
        meta_conf = round(qa_results.get("metadata_accuracy", 100.0) / 100.0, 2)
        
        final_overall_conf = round(qa_results.get("overall_score", 98.0) / 100.0, 2)
        review_required = final_overall_conf < 0.95 or qa_results.get("review_required", False)

        metadata_json["confidence_metrics"] = {
            "ocr_confidence": ocr_conf,
            "translation_confidence": trans_conf,
            "layout_confidence": layout_conf,
            "metadata_confidence": meta_conf,
            "overall_confidence": final_overall_conf,
            "review_required": review_required,
            "warnings": qa_results["warnings"]
        }

        doc_record = {
            "id": doc_id,
            "filename": display_name,
            "original_extension": ext,

            "language": primary_lang,
            "translated_language": "English",
            "upload_time": upload_time_str,
            "processing_time": processing_time,
            "status": "completed",
            "confidence": qa_results.get("overall_score", 98.0) / 100.0,
            "quality_score": qa_results.get("overall_score", 98.0) / 100.0,
            "quality_report": qa_results,
            "review_required": review_required,
            "original_pdf_path": orig_file_path,
            "translated_pdf_path": "",
            "translated_same_format_path": "",
            "metadata_json": metadata_json
        }

        # Step 13: Store record & extracted texts in FileDB
        await db_client.save_document(doc_record)
        await db_client.save_extracted_texts(doc_id, paragraphs)

        doc_record["paragraphs"] = paragraphs
        logger.info(f"Completed processing document {doc_id} ({clean_name}) in {processing_time}s")
        return doc_record

document_parser = DocumentParserService()
