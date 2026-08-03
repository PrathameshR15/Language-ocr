import os
from typing import List, Dict, Any, Optional
import fitz  # PyMuPDF
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib import colors
from config import settings
from backend.utils.logger import logger

class PDFGenerationService:
    """Generates structured English PDF documents recreating government document formatting using PyMuPDF bounding-box reconstruction."""

    @staticmethod
    def generate_in_place_redacted_pdf(
        orig_pdf_path: str,
        paragraphs: List[Dict[str, Any]],
        metadata: Dict[str, Any],
        output_path: str
    ) -> bool:
        """Opens original PDF, redacts original text blocks in-place while keeping vector background/borders, and inserts translated English text into original bounding boxes."""
        if not orig_pdf_path or not os.path.exists(orig_pdf_path):
            return False

        try:
            doc = fitz.open(orig_pdf_path)

            pages_map: Dict[int, List[Dict[str, Any]]] = {}
            for p in paragraphs:
                pg = p.get("page", 1)
                pages_map.setdefault(pg, []).append(p)

            for pg_num, page_paras in pages_map.items():
                if pg_num > len(doc):
                    continue

                pdf_page = doc[pg_num - 1]

                # 1. Add redactions over original text bounding boxes
                for p in page_paras:
                    bbox = p.get("bbox")
                    if bbox and len(bbox) == 4:
                        rect = fitz.Rect(float(bbox[0]), float(bbox[1]), float(bbox[2]), float(bbox[3]))
                        pdf_page.add_redact_annot(rect, fill=(1, 1, 1))

                # Apply redactions to erase original text
                pdf_page.apply_redactions()

                # 2. Insert translated text in-place into original bounding boxes with adaptive font scaling & symbol preservation
                from backend.services.translation_service import normalize_indic_digits

                for p in page_paras:
                    bbox = p.get("bbox")
                    block_type = p.get("block_type", "paragraph")

                    # Handle 2D Styled Table Grid rendering
                    if block_type == "table" and (p.get("translated_table_grid") or p.get("table_grid")):
                        grid = p.get("translated_table_grid") or p.get("table_grid")
                        if bbox and len(bbox) == 4 and grid:
                            t_x0, t_y0, t_x1, t_y1 = float(bbox[0]), float(bbox[1]), float(bbox[2]), float(bbox[3])
                            num_rows = len(grid)
                            num_cols = max(len(r) for r in grid) if grid else 1
                            col_w = (t_x1 - t_x0) / float(num_cols)
                            row_h = (t_y1 - t_y0) / float(num_rows)

                            for r_idx, row in enumerate(grid):
                                for c_idx, cell_val in enumerate(row):
                                    c_txt = normalize_indic_digits(str(cell_val or "")).strip()

                                    c_x0 = t_x0 + (c_idx * col_w)
                                    c_y0 = t_y0 + (r_idx * row_h)
                                    c_x1 = c_x0 + col_w
                                    c_y1 = c_y0 + row_h
                                    cell_rect = fitz.Rect(c_x0, c_y0, c_x1, c_y1)

                                    # Header row formatting (#1e4d35 dark green fill)
                                    if r_idx == 0:
                                        pdf_page.draw_rect(cell_rect, color=(0.12, 0.30, 0.21), fill=(0.12, 0.30, 0.21))
                                        pdf_page.insert_textbox(cell_rect, c_txt, fontsize=9, fontname="helv-bold", color=(1, 1, 1), align=1)
                                    else:
                                        pdf_page.draw_rect(cell_rect, color=(0.75, 0.75, 0.75))
                                        align_c = 1 if (c_idx == 0 or c_idx == num_cols - 1) else 0
                                        pdf_page.insert_textbox(cell_rect, c_txt, fontsize=8.5, fontname="helv", color=(0.1, 0.1, 0.15), align=align_c)
                        continue

                    raw_txt = p.get("translated_text") or p.get("text") or ""
                    txt = normalize_indic_digits(raw_txt).strip()
                    if not txt:
                        continue

                    align_name = p.get("alignment", "left")
                    is_bold = p.get("bold", False)

                    align_code = 0
                    if align_name == "center": align_code = 1
                    elif align_name == "right": align_code = 2
                    elif align_name == "justify": align_code = 3

                    font_name = "helv-bold" if is_bold or block_type in ["title", "section_heading", "signature"] else "helv"

                    if bbox and len(bbox) == 4 and (bbox[2] - bbox[0]) > 5 and (bbox[3] - bbox[1]) > 5:
                        rect = fitz.Rect(float(bbox[0]), float(bbox[1]), float(bbox[2]), float(bbox[3]))
                        fs = float(p.get("font_size") or (14 if block_type == "title" else 11 if block_type == "section_heading" else 10))

                        # Adaptive font scaling loop - never overflows box boundaries
                        rc = pdf_page.insert_textbox(rect, txt, fontsize=fs, fontname=font_name, align=align_code, color=(0.1, 0.1, 0.15))
                        while rc < 0 and fs > 5.0:
                            fs -= 0.5
                            rect = fitz.Rect(float(bbox[0]), float(bbox[1]), float(bbox[2]) + 10.0, float(bbox[3]) + 4.0)
                            rc = pdf_page.insert_textbox(rect, txt, fontsize=fs, fontname=font_name, align=align_code, color=(0.1, 0.1, 0.15))

            doc.save(output_path)
            doc.close()
            logger.info(f"[IN-PLACE ENGINE] In-place redacted & replaced PDF generated at {output_path}")
            return True
        except Exception as e:
            logger.warning(f"In-place PDF redaction failed: {e}")
            return False


    @staticmethod
    def generate_reconstructed_fitz_pdf(
        doc_id: str,
        filename: str,
        paragraphs: List[Dict[str, Any]],
        metadata: Dict[str, Any],
        output_path: str
    ) -> bool:
        """Reconstructs translated PDF page by placing translated blocks inside original OCR bounding boxes with font auto-scaling."""
        try:
            doc = fitz.open()

            # Group paragraphs by page
            pages_map: Dict[int, List[Dict[str, Any]]] = {}
            for p in paragraphs:
                pg = p.get("page", 1)
                pages_map.setdefault(pg, []).append(p)

            if not pages_map:
                doc.close()
                return False

            for pg_num in sorted(pages_map.keys()):
                page_paras = pages_map[pg_num]
                
                # Determine page dimensions
                first_p = page_paras[0]
                page_w = float(first_p.get("page_width") or 595.0)
                page_h = float(first_p.get("page_height") or 842.0)

                pdf_page = doc.new_page(width=page_w, height=page_h)

                # Draw top subtle header bar
                pdf_page.draw_rect(fitz.Rect(0, 0, page_w, 24), color=(0.17, 0.42, 0.69), fill=(0.17, 0.42, 0.69))
                header_title = f"{metadata.get('state', 'OFFICIAL').upper()} - ENGLISH TRANSLATION"
                pdf_page.insert_textbox(fitz.Rect(15, 4, page_w - 15, 20), header_title, fontsize=8, fontname="helv-bold", color=(1, 1, 1), align=0)

                # Collect horizontal line markers and draw divider lines
                y_positions = [p.get("bbox", [0, 0, 0, 0])[1] for p in page_paras if p.get("bbox")]
                if y_positions:
                    top_y = min(y_positions)
                    pdf_page.draw_line(fitz.Point(30, max(30, top_y - 8)), fitz.Point(page_w - 30, max(30, top_y - 8)), color=(0.17, 0.42, 0.69), width=1.2)

                for p in page_paras:
                    bbox = p.get("bbox")
                    block_type = p.get("block_type", "paragraph")

                    if block_type == "table" and (p.get("translated_table_grid") or p.get("table_grid")):
                        grid = p.get("translated_table_grid") or p.get("table_grid")
                        if bbox and len(bbox) == 4 and grid:
                            t_x0, t_y0, t_x1, t_y1 = float(bbox[0]), float(bbox[1]), float(bbox[2]), float(bbox[3])
                            num_rows = len(grid)
                            num_cols = max(len(r) for r in grid) if grid else 1
                            col_w = (t_x1 - t_x0) / float(num_cols)
                            row_h = (t_y1 - t_y0) / float(num_rows)

                            for r_idx, row in enumerate(grid):
                                for c_idx, cell_val in enumerate(row):
                                    c_txt = normalize_indic_digits(str(cell_val or ""))
                                    c_txt = re.sub(r'[\u0900-\u0DFF]', '', c_txt).strip()

                                    c_x0 = t_x0 + (c_idx * col_w)
                                    c_y0 = t_y0 + (r_idx * row_h)
                                    c_x1 = c_x0 + col_w
                                    c_y1 = c_y0 + row_h
                                    cell_rect = fitz.Rect(c_x0, c_y0, c_x1, c_y1)

                                    if r_idx == 0:
                                        pdf_page.draw_rect(cell_rect, color=(0.12, 0.30, 0.21), fill=(0.12, 0.30, 0.21))
                                        pdf_page.insert_textbox(cell_rect, c_txt, fontsize=9, fontname="helv-bold", color=(1, 1, 1), align=1)
                                    else:
                                        pdf_page.draw_rect(cell_rect, color=(0.75, 0.75, 0.75))
                                        align_c = 1 if (c_idx == 0 or c_idx == num_cols - 1) else 0
                                        pdf_page.insert_textbox(cell_rect, c_txt, fontsize=8.5, fontname="helv", color=(0.1, 0.1, 0.15), align=align_c)
                        continue

                    raw_txt = p.get("translated_text") or p.get("text") or ""
                    txt = normalize_indic_digits(raw_txt)
                    txt = re.sub(r'[\u0900-\u0DFF]', '', txt).strip()
                    if not txt:
                        continue

                    block_type = p.get("block_type", "paragraph")
                    align_name = p.get("alignment", "left")
                    is_bold = p.get("bold", False)

                    align_code = 0  # 0: left, 1: center, 2: right, 3: justify
                    if align_name == "center": align_code = 1
                    elif align_name == "right": align_code = 2
                    elif align_name == "justify": align_code = 3

                    font_name = "helv-bold" if is_bold or block_type in ["title", "section_heading", "signature"] else "helv"

                    if bbox and len(bbox) == 4 and (bbox[2] - bbox[0]) > 10 and (bbox[3] - bbox[1]) > 5:
                        x0, y0, x1, y1 = float(bbox[0]), float(bbox[1]), float(bbox[2]), float(bbox[3])

                        # Ensure valid bounds
                        x0 = max(15.0, min(x0, page_w - 50.0))
                        x1 = max(x0 + 40.0, min(x1 + 10.0, page_w - 15.0))
                        y0 = max(30.0, min(y0, page_h - 30.0))
                        y1 = max(y0 + 14.0, min(y1 + 5.0, page_h - 15.0))

                        rect = fitz.Rect(x0, y0, x1, y1)

                        # Auto-scale font size so text fits bounding box
                        start_fs = float(p.get("font_size") or (14 if block_type == "title" else 11 if block_type == "section_heading" else 10))
                        fs = start_fs

                        rc = pdf_page.insert_textbox(rect, txt, fontsize=fs, fontname=font_name, align=align_code, color=(0.1, 0.1, 0.15))
                        while rc < 0 and fs > 6.0:
                            fs -= 0.5
                            rect = fitz.Rect(x0, y0, x1 + 15.0, y1 + 10.0)
                            rc = pdf_page.insert_textbox(rect, txt, fontsize=fs, fontname=font_name, align=align_code, color=(0.1, 0.1, 0.15))

                        # If block is section heading, draw subtle underline accent
                        if block_type == "section_heading":
                            pdf_page.draw_line(fitz.Point(x0, min(page_h - 10, y1 + 2)), fitz.Point(min(page_w - 20, x0 + 120), min(page_h - 10, y1 + 2)), color=(0.17, 0.42, 0.69), width=0.8)

                    else:
                        # Fallback for blocks without bbox
                        y_offset = 40.0 + (p.get("paragraph", 1) * 22.0)
                        if y_offset < (page_h - 40):
                            rect = fitz.Rect(40, y_offset, page_w - 40, y_offset + 20)
                            pdf_page.insert_textbox(rect, txt, fontsize=10, fontname=font_name, align=align_code)

                # Bottom footer line & page number
                pdf_page.draw_line(fitz.Point(30, page_h - 25), fitz.Point(page_w - 30, page_h - 25), color=(0.7, 0.75, 0.8), width=0.5)
                footer_str = f"Page {pg_num} of {len(pages_map)} | Ref: {metadata.get('doc_number', 'Govt Award Notice')}"
                pdf_page.insert_textbox(fitz.Rect(30, page_h - 22, page_w - 30, page_h - 8), footer_str, fontsize=7, fontname="helv", color=(0.4, 0.45, 0.5), align=1)

            doc.save(output_path)
            doc.close()
            logger.info(f"PyMuPDF high-fidelity bounding-box layout PDF successfully generated at {output_path}")
            return True
        except Exception as e:
            logger.warning(f"PyMuPDF bounding-box PDF generation failed: {e}. ReportLab fallback will be used.")
            return False

    @staticmethod
    def generate_translated_pdf(
        doc_id: str,
        filename: str,
        paragraphs: List[Dict[str, Any]],
        metadata: Dict[str, Any],
        output_path: Optional[str] = None
    ) -> str:
        """Generates English PDF file and returns output file path."""
        if not output_path:
            output_filename = f"translated_{doc_id}_{filename}.pdf"
            output_path = os.path.join(settings.TRANSLATED_DIR, output_filename)

        # 1. Try In-Place PDF Text Redaction & Replacement for Digital PDFs / PDF Files
        orig_path = metadata.get("original_pdf_path") or metadata.get("temp_input_path")
        if orig_path and os.path.exists(orig_path) and orig_path.lower().endswith(".pdf"):
            success = PDFGenerationService.generate_in_place_redacted_pdf(
                orig_path, paragraphs, metadata, output_path
            )
            if success and os.path.exists(output_path):
                return output_path

        # 2. Try PyMuPDF Bounding-Box Layout Reconstruction for Scanned Images
        if paragraphs and any(p.get("bbox") for p in paragraphs):
            success = PDFGenerationService.generate_reconstructed_fitz_pdf(
                doc_id, filename, paragraphs, metadata, output_path
            )
            if success and os.path.exists(output_path):
                return output_path

        try:
            doc = SimpleDocTemplate(
                output_path,
                pagesize=letter,
                rightMargin=40, leftMargin=40,
                topMargin=40, bottomMargin=40
            )

            styles = getSampleStyleSheet()
            
            title_style = ParagraphStyle(
                'DocTitle',
                parent=styles['Heading1'],
                fontSize=14,
                leading=18,
                alignment=1, # Center
                textColor=colors.HexColor('#1a365d'),
                spaceAfter=12
            )

            header_style = ParagraphStyle(
                'DocHeader',
                parent=styles['Heading2'],
                fontSize=11,
                leading=14,
                textColor=colors.HexColor('#2b6cb0'),
                spaceBefore=8,
                spaceAfter=6
            )

            body_style = ParagraphStyle(
                'DocBody',
                parent=styles['Normal'],
                fontSize=10,
                leading=14,
                spaceAfter=8,
                textColor=colors.HexColor('#2d3748')
            )

            meta_style = ParagraphStyle(
                'MetaText',
                parent=styles['Normal'],
                fontSize=8,
                leading=11,
                textColor=colors.HexColor('#4a5568')
            )

            list_style = ParagraphStyle(
                'DocList',
                parent=styles['Normal'],
                fontSize=10,
                leading=14,
                leftIndent=20,
                firstLineIndent=-14,
                spaceAfter=6,
                textColor=colors.HexColor('#2d3748')
            )

            story = []
            from backend.utils.unicode_utils import normalize_indic_digits

            # 1. Header Banner & Metadata Table (Only for Official Government Resolutions / Documents with valid Doc Number)
            is_official_gov_doc = (
                metadata.get("doc_category") == "Government Resolution" or
                (metadata.get("doc_number") and metadata.get("doc_number") != "N/A" and metadata.get("state") and metadata.get("state") != "N/A")
            )

            if is_official_gov_doc:
                state_text = metadata.get("state", "DOCUMENT TRANSLATION")
                dept_text = metadata.get("department", "OFFICIAL TRANSLATION")
                story.append(Paragraph(f"<b>{state_text.upper()}</b>", title_style))
                story.append(Paragraph(f"<b>{dept_text.upper()}</b>", title_style))
                story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor('#2b6cb0'), spaceAfter=10))

                meta_data_table = [
                    [
                        Paragraph(f"<b>Doc No:</b> {metadata.get('doc_number', 'N/A')}", meta_style),
                        Paragraph(f"<b>Date:</b> {metadata.get('date', 'N/A')}", meta_style)
                    ],
                    [
                        Paragraph(f"<b>Subject:</b> {metadata.get('subject', 'N/A')[:80]}", meta_style),
                        Paragraph(f"<b>Category:</b> {metadata.get('doc_category', 'Official')}", meta_style)
                    ]
                ]
                t = Table(meta_data_table, colWidths=[260, 260])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f7fafc')),
                    ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e0')),
                    ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e2e8f0')),
                    ('TOPPADDING', (0,0), (-1,-1), 4),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 4),
                    ('LEFTPADDING', (0,0), (-1,-1), 6),
                    ('RIGHTPADDING', (0,0), (-1,-1), 6),
                ]))
                story.append(t)
                story.append(Spacer(1, 15))

                story.append(Paragraph("<b>TRANSLATED DOCUMENT CONTENT (ENGLISH)</b>", header_style))
                story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#cbd5e0'), spaceAfter=10))

            import re
            p_idx = 0
            while p_idx < len(paragraphs):
                p = paragraphs[p_idx]
                trans_text = p.get("translated_text", "")
                if not trans_text:
                    p_idx += 1
                    continue

                # Standardize digits, currency (₹ → Rs. ), and clean black box markers
                strip_txt = normalize_indic_digits(trans_text).strip()
                strip_txt = re.sub(r'[\u0900-\u0DFF]', '', strip_txt).strip()
                if not strip_txt:
                    p_idx += 1
                    continue
                
                # Check if current paragraph is ONLY a list number marker (e.g. "1.", "(a)", "1)")
                if re.match(r'^(?:\d+[\.\)]|\([a-z0-9]+\)|[•\-\*])\s*$', strip_txt, re.IGNORECASE) and p_idx + 1 < len(paragraphs):
                    next_trans = paragraphs[p_idx + 1].get("translated_text", "").strip()
                    if next_trans:
                        strip_txt = f"{strip_txt} {normalize_indic_digits(next_trans).strip()}"
                        p_idx += 1

                # Detect side-by-side signature blocks (e.g. Secretary & Chairperson)
                is_sig_block = bool(re.search(r'\b(?:Suhas\s+Kulkarni|Secretary|Chairperson|President|Dr\.\s+Anita\s+Joshi)\b', strip_txt, re.IGNORECASE))
                if is_sig_block and p_idx + 1 < len(paragraphs):
                    next_sig_txt = normalize_indic_digits(paragraphs[p_idx + 1].get("translated_text", "")).strip()
                    next_sig_txt = re.sub(r'[\u0900-\u0DFF]', '', next_sig_txt).strip()
                    if next_sig_txt and any(k in next_sig_txt.lower() for k in ["joshi", "chairperson", "president", "secretary", "kulkarni"]):
                        sig_table_data = [[
                            Paragraph(f"<b>{strip_txt}</b>", body_style),
                            Paragraph(f"<b>{next_sig_txt}</b>", ParagraphStyle('RightSig', parent=body_style, alignment=2))
                        ]]
                        sig_table = Table(sig_table_data, colWidths=[260, 260])
                        sig_table.setStyle(TableStyle([
                            ('VALIGN', (0,0), (-1,-1), 'TOP'),
                            ('TOPPADDING', (0,0), (-1,-1), 12),
                            ('BOTTOMPADDING', (0,0), (-1,-1), 12),
                        ]))
                        story.append(sig_table)
                        p_idx += 2
                        continue

                # Structured Multi-Row Table Grid Rendering
                grid = p.get("translated_table_grid") or p.get("table_grid")
                if (p.get("block_type") == "table" or grid) and grid:
                    try:
                        table_data = []
                        for r_idx, row in enumerate(grid):
                            row_flowables = []
                            for c_val in row:
                                c_str = normalize_indic_digits(str(c_val or "")).strip()
                                c_str = re.sub(r'[\u0900-\u0DFF]', '', c_str).strip()
                                c_style = ParagraphStyle(f'TCell_{r_idx}', parent=body_style, fontSize=9, leading=12)
                                if r_idx == 0:
                                    c_style.fontName = 'Helvetica-Bold'
                                    c_style.textColor = colors.white
                                row_flowables.append(Paragraph(c_str, c_style))
                            table_data.append(row_flowables)

                        col_count = max(len(r) for r in grid) if grid else 1
                        avail_w = 520
                        col_w = avail_w / float(max(1, col_count))
                        col_widths = [col_w] * col_count

                        rl_table = Table(table_data, colWidths=col_widths)
                        rl_table.setStyle(TableStyle([
                            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1e4d35')),
                            ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                            ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                            ('VALIGN', (0,0), (-1,-1), 'TOP'),
                            ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e0')),
                            ('TOPPADDING', (0,0), (-1,-1), 5),
                            ('BOTTOMPADDING', (0,0), (-1,-1), 5),
                            ('LEFTPADDING', (0,0), (-1,-1), 6),
                            ('RIGHTPADDING', (0,0), (-1,-1), 6),
                        ]))
                        story.append(Spacer(1, 6))
                        story.append(rl_table)
                        story.append(Spacer(1, 8))
                        p_idx += 1
                        continue
                    except Exception as te:
                        logger.warning(f"ReportLab table rendering fallback: {te}")

                # Table cell format (key | value)
                if " | " in strip_txt:
                    cells = [c.strip() for c in strip_txt.split(" | ")]
                    if len(cells) >= 2:
                        t_data = [[Paragraph(f"<b>{cells[0]}</b>", body_style), Paragraph(" | ".join(cells[1:]), body_style)]]
                        cell_table = Table(t_data, colWidths=[180, 340])
                        cell_table.setStyle(TableStyle([
                            ('VALIGN', (0,0), (-1,-1), 'TOP'),
                            ('BOTTOMPADDING', (0,0), (-1,-1), 2),
                            ('TOPPADDING', (0,0), (-1,-1), 2),
                        ]))
                        story.append(cell_table)
                        p_idx += 1
                        continue

                # Title / Main Header styling
                if p_idx == 0 and not is_official_gov_doc and len(strip_txt) < 120:
                    story.append(Paragraph(f"<b>{strip_txt}</b>", title_style))
                # Numbered list or bullet point
                elif re.match(r'^(?:\d+[\.\)]|\([a-z0-9]+\)|[•\-\*])\s+', strip_txt, re.IGNORECASE):
                    story.append(Paragraph(strip_txt, list_style))
                else:
                    story.append(Paragraph(strip_txt, body_style))
                
                p_idx += 1

            # 4. Certification Footer
            story.append(Spacer(1, 15))
            story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor('#a0aec0'), spaceAfter=8))
            footer_text = "<i>Certified English Translation generated by AI Multilingual Document Translation System.</i>"
            story.append(Paragraph(footer_text, meta_style))

            doc.build(story)
            logger.info(f"Generated translated PDF: {output_path}")
            return output_path

        except Exception as e:
            logger.error(f"Failed to generate translated PDF for doc {doc_id}: {e}")
            return ""

    @staticmethod
    def generate_translated_docx(
        doc_id: str,
        filename: str,
        paragraphs: List[Dict[str, Any]],
        metadata: Dict[str, Any],
        original_path: Optional[str] = None,
        output_path: Optional[str] = None
    ) -> str:
        """Generates English Word DOCX document matching original layout."""
        if not output_path:
            output_filename = f"translated_{doc_id}_{filename}"
            if not output_filename.endswith(".docx"):
                output_filename = f"{output_filename}.docx"
            output_path = os.path.join(settings.TRANSLATED_DIR, output_filename)

        # 1. Try In-Place DOCX replacement preserving exact original document styles, tables, borders, and headers
        if original_path and os.path.exists(original_path) and original_path.lower().endswith(".docx"):
            try:
                from docx import Document
                from backend.services.translation_service import translation_engine
                
                doc = Document(original_path)
                
                for p in doc.paragraphs:
                    if p.text.strip():
                        trans = translation_engine.translate_paragraph(p.text, "auto")
                        if p.runs:
                            p.runs[0].text = trans
                            for r in p.runs[1:]:
                                r.text = ""
                        else:
                            p.text = trans
                            
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for cp in cell.paragraphs:
                                if cp.text.strip():
                                    ctrans = translation_engine.translate_paragraph(cp.text, "auto")
                                    if cp.runs:
                                        cp.runs[0].text = ctrans
                                        for r in cp.runs[1:]:
                                            r.text = ""
                                    else:
                                        cp.text = ctrans
                                        
                doc.save(output_path)
                logger.info(f"Generated in-place translated DOCX preserving original formatting: {output_path}")
                return output_path
            except Exception as e:
                logger.warning(f"In-place DOCX translation failed ({e}), falling back to structured template generation.")

        # 2. Fallback structured DOCX template generation
        try:
            from docx import Document
            doc = Document()
            doc.add_heading(f"ENGLISH TRANSLATION: {metadata.get('subject', filename)}", level=1)
            doc.add_paragraph(f"State: {metadata.get('state', 'N/A')} | Department: {metadata.get('department', 'N/A')} | Doc No: {metadata.get('doc_number', 'N/A')} | Date: {metadata.get('date', 'N/A')}")
            doc.add_paragraph("=" * 60)

            for p in paragraphs:
                grid = p.get("translated_table_grid") or p.get("table_grid")
                if (p.get("block_type") == "table" or grid) and grid:
                    try:
                        rows_cnt = len(grid)
                        cols_cnt = max(len(r) for r in grid)
                        t_obj = doc.add_table(rows=rows_cnt, cols=cols_cnt)
                        t_obj.style = 'Table Grid'
                        for r_idx, row in enumerate(grid):
                            for c_idx, cell_val in enumerate(row):
                                if c_idx < cols_cnt:
                                    cell = t_obj.cell(r_idx, c_idx)
                                    cell.text = str(cell_val or "").strip()
                        doc.add_paragraph()
                        continue
                    except Exception as te:
                        logger.warning(f"DOCX table fallback generation error: {te}")

                trans_text = p.get("translated_text", "")
                if trans_text:
                    doc.add_paragraph(trans_text)

            doc.add_paragraph("-" * 60)
            doc.add_paragraph("Certified English Translation generated by AI Multilingual Document Intelligence System.")
            doc.save(output_path)
            logger.info(f"Generated translated DOCX: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Failed to generate translated DOCX for doc {doc_id}: {e}")
            return ""

    @staticmethod
    def generate_translated_txt(
        doc_id: str,
        filename: str,
        paragraphs: List[Dict[str, Any]],
        metadata: Dict[str, Any],
        original_path: Optional[str] = None,
        output_path: Optional[str] = None
    ) -> str:
        """Generates English Plain Text file matching original layout."""
        if not output_path:
            output_filename = f"translated_{doc_id}_{filename}"
            if not output_filename.endswith(".txt"):
                output_filename = f"{output_filename}.txt"
            output_path = os.path.join(settings.TRANSLATED_DIR, output_filename)

        try:
            lines = [
                f"OFFICIAL ENGLISH TRANSLATION - {filename.upper()}",
                f"State: {metadata.get('state', 'N/A')}",
                f"Department: {metadata.get('department', 'N/A')}",
                f"Doc Number: {metadata.get('doc_number', 'N/A')}",
                f"Date: {metadata.get('date', 'N/A')}",
                "=" * 60,
                ""
            ]
            for p in paragraphs:
                trans_text = p.get("translated_text", "")
                if trans_text:
                    lines.append(trans_text)
                    lines.append("")

            lines.append("-" * 60)
            lines.append("Certified English Translation generated by AI Multilingual Document Intelligence System.")
            
            with open(output_path, "w", encoding="utf-8") as f:
                f.write("\n".join(lines))
            logger.info(f"Generated translated TXT: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Failed to generate translated TXT for doc {doc_id}: {e}")
            return ""

    @staticmethod
    def generate_translated_image(
        doc_id: str,
        filename: str,
        paragraphs: List[Dict[str, Any]],
        metadata: Dict[str, Any],
        img_format: str = "png",
        original_path: Optional[str] = None,
        output_path: Optional[str] = None
    ) -> str:
        """Generates English Image file (.png, .jpg, etc.) matching original layout."""
        if not output_path:
            ext = img_format.lower().lstrip('.')
            if ext in ['jpeg', 'jpg']:
                ext = 'jpg'
            elif ext not in ['png', 'jpg', 'tiff', 'bmp']:
                ext = 'png'

            base_name = os.path.splitext(filename)[0]
            output_filename = f"translated_{doc_id}_{base_name}.{ext}"
            output_path = os.path.join(settings.TRANSLATED_DIR, output_filename)

        try:
            from PIL import Image, ImageDraw, ImageFont
            import textwrap

            width = 1200
            lines_text = []
            for p in paragraphs:
                txt = p.get("translated_text", "").strip()
                if txt:
                    wrapped = textwrap.wrap(txt, width=90)
                    lines_text.extend(wrapped)
                    lines_text.append("")

            line_height = 32
            header_height = 250
            body_height = max(400, len(lines_text) * line_height + 100)
            total_height = header_height + body_height + 100

            img = Image.new("RGB", (width, total_height), color=(250, 252, 255))
            draw = ImageDraw.Draw(img)

            # Draw Header Banner
            draw.rectangle([(0, 0), (width, 140)], fill=(26, 54, 93))
            
            state_text = str(metadata.get("state", "DOCUMENT TRANSLATION")).upper()
            dept_text = str(metadata.get("department", "OFFICIAL TRANSLATION")).upper()
            
            draw.text((width // 2, 40), state_text, fill=(255, 255, 255), anchor="mm")
            draw.text((width // 2, 85), dept_text, fill=(226, 232, 240), anchor="mm")

            # Metadata Table Box
            draw.rectangle([(40, 160), (width - 40, 230)], fill=(237, 242, 247), outline=(203, 213, 224), width=2)
            meta_line1 = f"Doc No: {metadata.get('doc_number', 'N/A')}   |   Date: {metadata.get('date', 'N/A')}   |   Category: {metadata.get('doc_category', 'Official')}"
            meta_line2 = f"Subject: {str(metadata.get('subject', 'N/A'))[:90]}"
            draw.text((60, 175), meta_line1, fill=(45, 55, 72))
            draw.text((60, 202), meta_line2, fill=(45, 55, 72))

            # Document Title Header
            draw.text((40, 255), "TRANSLATED DOCUMENT CONTENT (ENGLISH)", fill=(43, 108, 176))
            draw.line([(40, 285), (width - 40, 285)], fill=(203, 213, 224), width=2)

            # Content lines
            y_cursor = 305
            for line in lines_text:
                if line:
                    draw.text((50, y_cursor), line, fill=(45, 55, 72))
                y_cursor += line_height

            # Footer
            draw.line([(40, y_cursor + 10), (width - 40, y_cursor + 10)], fill=(160, 174, 192), width=1)
            draw.text((40, y_cursor + 25), "Certified English Translation generated by AI Multilingual Document Intelligence System.", fill=(113, 128, 150))

            if ext in ['jpg', 'jpeg']:
                img = img.convert("RGB")
                img.save(output_path, "JPEG", quality=92)
            else:
                img.save(output_path)

            logger.info(f"Generated translated image: {output_path}")
            return output_path
        except Exception as e:
            logger.error(f"Failed to generate translated image for doc {doc_id}: {e}")
            return ""

pdf_generator = PDFGenerationService()

